"""
文献导出模块 - 将论文列表导出为 BibTeX / CSV / Excel / Word 格式

功能:
  1. 导出 BibTeX 格式（.bib）- 可直接导入 LaTeX
  2. 导出 CSV 格式（.csv）- 可用 Excel 打开
  3. 导出 Excel 格式（.xlsx）- 带格式的 Excel 表格
  4. 导出 Word 综述文档（.docx）- 带格式的综述 + 参考文献
  5. 文件保存到 exports/ 目录，通过 Web 服务提供下载
"""

import os
import csv
import re
import logging
from datetime import datetime
from typing import List, Optional

from .fetcher import Paper

logger = logging.getLogger(__name__)

EXPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")


class LiteratureExporter:
    """文献导出器"""

    def __init__(self, config: dict = None):
        self.config = config or {}
        os.makedirs(EXPORTS_DIR, exist_ok=True)

    def export_bibtex(self, papers: List[Paper], filename_prefix: str = "papers") -> Optional[str]:
        """
        导出论文列表为 BibTeX 格式

        Args:
            papers: 论文列表
            filename_prefix: 文件名前缀

        Returns:
            文件绝对路径，失败返回 None
        """
        if not papers:
            return None

        entries = []
        for i, p in enumerate(papers, 1):
            entry = self._paper_to_bibtex(p, i)
            entries.append(entry)

        content = "\n\n".join(entries)

        filepath = os.path.join(EXPORTS_DIR, f"{filename_prefix}.bib")
        try:
            with open(filepath, "w", encoding="utf-8") as f:
                f.write(content)
            logger.info(f"BibTeX 导出成功: {filepath} ({len(papers)} 篇)")
            return filepath
        except Exception as e:
            logger.error(f"BibTeX 导出失败: {e}")
            return None

    def export_csv(self, papers: List[Paper], filename_prefix: str = "papers") -> Optional[str]:
        """
        导出论文列表为 CSV 格式

        Args:
            papers: 论文列表
            filename_prefix: 文件名前缀

        Returns:
            文件绝对路径
        """
        if not papers:
            return None

        filepath = os.path.join(EXPORTS_DIR, f"{filename_prefix}.csv")
        try:
            with open(filepath, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                # 表头
                writer.writerow([
                    "序号", "标题", "作者", "年份", "期刊/会议",
                    "引用数", "DOI", "URL", "摘要", "关键词"
                ])
                for i, p in enumerate(papers, 1):
                    year = getattr(p, '_year', '') or p.published.year if p.published else ''
                    venue = getattr(p, '_venue', '') or ''
                    citation = getattr(p, '_citation_count', 0) or 0
                    doi = getattr(p, '_doi', '') or ''
                    authors = '; '.join(p.authors[:10])

                    writer.writerow([
                        i,
                        p.title,
                        authors,
                        year,
                        venue,
                        citation,
                        doi,
                        p.url,
                        p.abstract[:300] if p.abstract else '',
                        ', '.join(p.keywords_matched) if p.keywords_matched else '',
                    ])

            logger.info(f"CSV 导出成功: {filepath} ({len(papers)} 篇)")
            return filepath
        except Exception as e:
            logger.error(f"CSV 导出失败: {e}")
            return None

    def export_excel(self, papers: List[Paper], filename_prefix: str = "papers") -> Optional[str]:
        """
        导出论文列表为 Excel 格式（.xlsx），带格式化

        Args:
            papers: 论文列表
            filename_prefix: 文件名前缀

        Returns:
            文件绝对路径
        """
        if not papers:
            return None

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        except ImportError:
            logger.error("openpyxl 未安装，无法导出 Excel。运行: pip install openpyxl")
            return None

        filepath = os.path.join(EXPORTS_DIR, f"{filename_prefix}.xlsx")
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "文献列表"

            # 表头样式
            header_font = Font(bold=True, color="FFFFFF", size=11)
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            thin_border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            # 表头
            headers = ["序号", "标题", "作者", "年份", "期刊/会议", "引用数", "DOI", "URL", "摘要"]
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=1, column=col, value=header)
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
                cell.border = thin_border

            # 数据行
            for i, p in enumerate(papers, 1):
                year = getattr(p, '_year', '') or ''
                venue = getattr(p, '_venue', '') or ''
                citation = getattr(p, '_citation_count', 0) or 0
                doi = getattr(p, '_doi', '') or ''
                authors = '; '.join(p.authors[:10])
                abstract = p.abstract[:500] if p.abstract else ''

                row_data = [i, p.title, authors, year, venue, citation, doi, p.url, abstract]
                for col, value in enumerate(row_data, 1):
                    cell = ws.cell(row=i + 1, column=col, value=value)
                    cell.border = thin_border
                    cell.alignment = Alignment(vertical="top", wrap_text=(col in [2, 3, 9]))

            # 设置列宽
            col_widths = [6, 60, 30, 8, 25, 10, 20, 40, 60]
            for col, width in enumerate(col_widths, 1):
                ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = width

            # 冻结首行
            ws.freeze_panes = "A2"

            wb.save(filepath)
            logger.info(f"Excel 导出成功: {filepath} ({len(papers)} 篇)")
            return filepath
        except Exception as e:
            logger.error(f"Excel 导出失败: {e}")
            return None

    def export_review_docx(self, topic: str, review_text: str,
                           papers: List[Paper],
                           filename_prefix: str = "review") -> Optional[str]:
        """
        导出文献综述为 Word (.docx) 文档，包含综述正文和参考文献列表

        Args:
            topic: 综述主题
            review_text: AI 生成的综述文本（Markdown 格式）
            papers: 参考文献列表
            filename_prefix: 文件名前缀

        Returns:
            文件绝对路径
        """
        if not review_text:
            return None

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
        except ImportError:
            logger.error("python-docx 未安装，无法导出 Word。运行: pip install python-docx")
            return None

        filepath = os.path.join(EXPORTS_DIR, f"{filename_prefix}.docx")
        try:
            doc = Document()

            # ---- 设置默认字体 ----
            style = doc.styles['Normal']
            font = style.font
            font.name = '宋体'
            font.size = Pt(10.5)

            # ---- 页面标题 ----
            title_para = doc.add_heading(f'文献综述：{topic}', level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 基本信息
            from datetime import datetime as dt
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = info_para.add_run(f'生成日期：{dt.now().strftime("%Y年%m月%d日")}  |  '
                                    f'参考文献：{len(papers)} 篇')
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)

            doc.add_paragraph()  # 空行

            # ---- 解析 Markdown 综述文本写入 Word ----
            self._write_markdown_to_docx(doc, review_text)

            # ---- 参考文献部分 ----
            doc.add_page_break()
            doc.add_heading('参考文献', level=1)

            for i, p in enumerate(papers, 1):
                year = getattr(p, '_year', '') or ''
                venue = getattr(p, '_venue', '') or ''
                citation = getattr(p, '_citation_count', 0) or 0
                doi = getattr(p, '_doi', '') or ''
                authors = ', '.join(p.authors[:5])
                if len(p.authors) > 5:
                    authors += ' et al.'

                # 参考文献格式: [序号] 作者. 标题. 期刊, 年份. DOI
                ref_para = doc.add_paragraph(style='List Number')
                ref_run = ref_para.add_run(f'[{i}] {authors}. ')
                ref_run.font.size = Pt(9)

                # 标题（斜体）
                title_run = ref_para.add_run(f'{p.title}. ')
                title_run.font.size = Pt(9)
                title_run.italic = True

                # 期刊/年份/引用
                meta_parts = []
                if venue:
                    meta_parts.append(venue)
                if year:
                    meta_parts.append(str(year))
                if doi:
                    meta_parts.append(f'DOI: {doi}')
                meta_parts.append(f'引用: {citation}')
                meta_run = ref_para.add_run('. '.join(meta_parts))
                meta_run.font.size = Pt(9)
                meta_run.font.color.rgb = RGBColor(80, 80, 80)

                # URL
                if p.url:
                    url_run = ref_para.add_run(f'\n    {p.url}')
                    url_run.font.size = Pt(8)
                    url_run.font.color.rgb = RGBColor(0, 102, 204)

            # ---- 保存 ----
            doc.save(filepath)
            logger.info(f"Word 综述导出成功: {filepath} ({len(papers)} 篇参考文献)")
            return filepath
        except Exception as e:
            logger.error(f"Word 综述导出失败: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _write_markdown_to_docx(self, doc, text: str):
        """
        将 Markdown 格式文本解析写入 Word 文档

        支持:
        - # / ## / ### 标题
        - **粗体** 和 *斜体*
        - 列表项 (- 和 数字.)
        - 普通段落
        """
        from docx.shared import Pt, RGBColor

        lines = text.split('\n')
        for line in lines:
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph()  # 空行
                continue

            # 标题
            if stripped.startswith('### '):
                doc.add_heading(stripped[4:].strip(), level=3)
            elif stripped.startswith('## '):
                doc.add_heading(stripped[3:].strip(), level=2)
            elif stripped.startswith('# '):
                doc.add_heading(stripped[2:].strip(), level=1)
            # 列表项
            elif re.match(r'^[-*]\s+', stripped):
                content = re.sub(r'^[-*]\s+', '', stripped)
                para = doc.add_paragraph(style='List Bullet')
                self._add_formatted_runs(para, content)
            elif re.match(r'^\d+\.\s+', stripped):
                content = re.sub(r'^\d+\.\s+', '', stripped)
                para = doc.add_paragraph(style='List Number')
                self._add_formatted_runs(para, content)
            else:
                # 普通段落
                para = doc.add_paragraph()
                self._add_formatted_runs(para, stripped)

    def _add_formatted_runs(self, paragraph, text: str):
        """
        解析 Markdown 内联格式并添加到段落

        支持 **粗体**, *斜体*, [序号] 引用标记
        """
        from docx.shared import Pt, RGBColor

        # 用正则拆分文本: **粗体**、*斜体*、[n] 引用
        pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|\[\d+\])'
        parts = re.split(pattern, text)

        for part in parts:
            if not part:
                continue
            if part.startswith('**') and part.endswith('**'):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
                run.font.size = Pt(10.5)
            elif part.startswith('*') and part.endswith('*'):
                run = paragraph.add_run(part[1:-1])
                run.italic = True
                run.font.size = Pt(10.5)
            elif re.match(r'^\[\d+\]$', part):
                # 引用标记 [1], [2] 等用上标蓝色
                run = paragraph.add_run(part)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0, 102, 204)
                run.font.superscript = True
            else:
                run = paragraph.add_run(part)
                run.font.size = Pt(10.5)

    def export_markdown_table(self, papers: List[Paper]) -> str:
        """
        生成 Markdown 表格（直接返回文本，可在钉钉显示）

        Returns:
            Markdown 格式的表格文本
        """
        lines = []
        lines.append("| # | 标题 | 作者 | 年份 | 引用 | 链接 |")
        lines.append("|---|------|------|------|------|------|")

        for i, p in enumerate(papers, 1):
            year = getattr(p, '_year', '') or ''
            citation = getattr(p, '_citation_count', 0) or 0
            authors = ', '.join(p.authors[:2])
            if len(p.authors) > 2:
                authors += ' et al.'
            title_short = p.title[:50] + ('...' if len(p.title) > 50 else '')

            lines.append(f"| {i} | {title_short} | {authors} | {year} | {citation} | [链接]({p.url}) |")

        return "\n".join(lines)

    def _paper_to_bibtex(self, paper: Paper, index: int) -> str:
        """将 Paper 对象转为 BibTeX 条目"""
        # 生成引用key
        first_author = paper.authors[0] if paper.authors else "Unknown"
        last_name = first_author.split()[-1] if first_author else "Unknown"
        year = getattr(paper, '_year', '') or (paper.published.year if paper.published else datetime.now().year)
        # 清理特殊字符
        clean_name = re.sub(r'[^a-zA-Z]', '', last_name)
        cite_key = f"{clean_name}{year}_{index}"

        # 提取元数据
        doi = getattr(paper, '_doi', '') or ''
        arxiv_raw = getattr(paper, '_arxiv_id_raw', '') or ''
        venue = getattr(paper, '_venue', '') or ''

        # 格式化作者列表（BibTeX 格式: Last, First and Last, First）
        bib_authors = []
        for author in paper.authors[:15]:
            parts = author.strip().split()
            if len(parts) >= 2:
                bib_authors.append(f"{parts[-1]}, {' '.join(parts[:-1])}")
            else:
                bib_authors.append(author)
        authors_str = " and ".join(bib_authors)

        # 转义 BibTeX 特殊字符
        title_escaped = paper.title.replace('{', '\\{').replace('}', '\\}')
        abstract_escaped = (paper.abstract[:500] if paper.abstract else '').replace('{', '\\{').replace('}', '\\}')

        # 使用 @article 还是 @inproceedings
        entry_type = "article"
        if any(kw in (venue or '').lower() for kw in ['conference', 'proceedings', 'workshop', 'symposium']):
            entry_type = "inproceedings"

        lines = [f"@{entry_type}{{{cite_key},"]
        lines.append(f"  title = {{{title_escaped}}},")
        lines.append(f"  author = {{{authors_str}}},")
        lines.append(f"  year = {{{year}}},")
        if venue:
            if entry_type == "article":
                lines.append(f"  journal = {{{venue}}},")
            else:
                lines.append(f"  booktitle = {{{venue}}},")
        if doi:
            lines.append(f"  doi = {{{doi}}},")
        if arxiv_raw:
            lines.append(f"  eprint = {{{arxiv_raw}}},")
            lines.append(f"  archivePrefix = {{arXiv}},")
        lines.append(f"  url = {{{paper.url}}},")
        if abstract_escaped:
            lines.append(f"  abstract = {{{abstract_escaped}}},")
        lines.append("}")

        return "\n".join(lines)
