#!/usr/bin/env python3
"""
钉钉 AI 助理机器人 - 基于 Stream 模式（长连接）

功能:
  1. 接收钉钉群内 @机器人 的消息
  2. 调用 AI 研究助理处理用户请求
  3. 返回文献搜索、综述、导出等结果
  4. 提供文件下载服务（BibTeX/CSV）

前置条件:
  1. 在钉钉开放平台创建「企业内部应用」
  2. 启用「机器人」功能，选择 Stream 模式
  3. 将 AppKey 和 AppSecret 填入 config.yaml

使用方式:
  python dingtalk_bot_server.py                     # 启动机器人
  python dingtalk_bot_server.py --port 5679         # 指定文件服务端口
  python dingtalk_bot_server.py --webhook-only       # 仅使用 Webhook 推送模式（不需要企业应用）
"""

import os
import sys
import argparse
import logging
import asyncio
import json
import threading
from pathlib import Path

sys.path.insert(0, os.path.dirname(__file__))

from flask import Flask, send_from_directory, jsonify
from src.utils import load_config
from src.research_assistant import ResearchAssistant

logger = logging.getLogger(__name__)
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s"
)

# 全局配置
EXPORTS_DIR = os.path.join(os.path.dirname(__file__), "exports")

# ============================================================
# Flask 文件下载服务
# ============================================================

file_app = Flask(__name__)


@file_app.route("/download/<filename>")
def download_file(filename):
    """提供导出文件的下载服务（exports + chat_logs）"""
    safe_name = os.path.basename(filename)
    # 先在 exports 目录中找
    filepath = os.path.join(EXPORTS_DIR, safe_name)
    if os.path.exists(filepath):
        return send_from_directory(EXPORTS_DIR, safe_name, as_attachment=True)
    # 再在 chat_logs 目录中找
    chat_logs_dir = os.path.join(os.path.dirname(__file__), "chat_logs")
    filepath2 = os.path.join(chat_logs_dir, safe_name)
    if os.path.exists(filepath2):
        return send_from_directory(chat_logs_dir, safe_name, as_attachment=True)
    return jsonify({"error": "文件不存在"}), 404


@file_app.route("/exports")
def list_exports():
    """列出可下载的文件"""
    os.makedirs(EXPORTS_DIR, exist_ok=True)
    files = []
    for f in sorted(os.listdir(EXPORTS_DIR), reverse=True):
        filepath = os.path.join(EXPORTS_DIR, f)
        if os.path.isfile(filepath):
            size = os.path.getsize(filepath)
            files.append({"name": f, "size": size, "url": f"/download/{f}"})
    return jsonify(files)


def start_file_server(port: int):
    """在后台线程中启动文件下载服务"""
    file_app.run(host="0.0.0.0", port=port, debug=False, use_reloader=False)


# ============================================================
# 钉钉 Stream 模式机器人
# ============================================================

class DingTalkBotStream:
    """
    钉钉 Stream 模式机器人

    使用长连接（WebSocket）与钉钉服务器通信，
    不需要公网IP或域名，适合本地/内网部署。
    """

    def __init__(self, config: dict):
        self.config = config
        bot_config = config.get("dingtalk_bot", {})
        self.app_key = bot_config.get("app_key", "")
        self.app_secret = bot_config.get("app_secret", "")
        self.file_server_url = bot_config.get("file_server_url", "http://127.0.0.1:5679")

        if not self.app_key or not self.app_secret:
            raise ValueError(
                "钉钉 AI 助理需要企业内部应用的 AppKey 和 AppSecret。\n"
                "请在 config.yaml 的 dingtalk_bot 部分填入。\n"
                "创建应用: https://open-dev.dingtalk.com → 应用开发 → 企业内部应用"
            )

        self.assistant = ResearchAssistant(config)

    def start(self):
        """启动 Stream 模式机器人"""
        try:
            import dingtalk_stream
            from dingtalk_stream import AckMessage
        except ImportError:
            logger.error(
                "dingtalk-stream 未安装。请运行: pip install dingtalk-stream\n"
                "文档: https://github.com/open-dingtalk/dingtalk-stream-sdk-python"
            )
            sys.exit(1)

        # 启用 WebSocket 和 SDK 的详细日志
        logging.getLogger('dingtalk_stream').setLevel(logging.DEBUG)
        logging.getLogger('websockets').setLevel(logging.DEBUG)

        credential = dingtalk_stream.Credential(self.app_key, self.app_secret)
        client = dingtalk_stream.DingTalkStreamClient(credential)

        # 创建聊天机器人消息处理器（继承 SDK 的 ChatbotHandler）
        handler = MyChatbotHandler(self.assistant, self.config)

        # 注册消息回调
        client.register_callback_handler(
            dingtalk_stream.ChatbotMessage.TOPIC,
            handler
        )

        logger.info("🤖 钉钉 AI 助理启动中（Stream 模式）...")
        logger.info(f"   AppKey: {self.app_key[:8]}...")
        logger.info(f"   Topic: {dingtalk_stream.ChatbotMessage.TOPIC}")
        client.start_forever()


class MyChatbotHandler:
    """
    聊天机器人消息处理器

    继承 dingtalk_stream.ChatbotHandler 以使用内置的 reply_text / reply_markdown 方法。
    注意: 实际继承在 _create_handler_class() 中动态完成（避免顶层 import）。
    这里先定义为一个代理类，在 __init__ 中动态转换。
    """

    def __new__(cls, assistant, config):
        """动态创建继承自 dingtalk_stream.ChatbotHandler 的实例"""
        import dingtalk_stream

        # 动态创建子类
        class _Handler(dingtalk_stream.ChatbotHandler):
            def __init__(self, _assistant, _config):
                super().__init__()
                self._assistant = _assistant
                self._config = _config

            async def process(self, callback):
                from dingtalk_stream import AckMessage, ChatbotMessage
                import time as _time
                incoming_message = ChatbotMessage.from_dict(callback.data)

                # 提取文本内容
                text_list = self.extract_text_from_incoming_message(incoming_message)
                user_message = " ".join(text_list).strip() if text_list else ""
                user_id = incoming_message.sender_id or ""
                sender_nick = getattr(incoming_message, 'sender_nick', '') or ''

                logger.info(f"收到钉钉消息 [{sender_nick}({user_id})]: {user_message[:100]}")

                if not user_message:
                    self.reply_text("请输入您的问题或指令。发送「帮助」查看使用指南。",
                                    incoming_message)
                    return AckMessage.STATUS_OK, "OK"

                # 异步后台处理，立即返回 ACK，避免超时
                async def _background_process():
                    try:
                        # 先发送"处理中"提示
                        logger.info("[DEBUG] 发送'处理中'提示...")
                        self.reply_text("🔍 正在处理您的请求，请稍候...", incoming_message)
                        logger.info("[DEBUG] '处理中'提示已发送")
    
                        # 调用 AI 研究助理
                        # handle_message 是 async def 但内部全是同步调用
                        # 用 run_in_executor 在线程池中运行，避免阻塞 WebSocket 事件循环
                        logger.info("[DEBUG] 开始调用 handle_message...")
                        _t0 = _time.time()
                        import functools
                        loop = asyncio.get_event_loop()
    
                        def _sync_handle():
                            """在新事件循环中运行 async handle_message"""
                            _loop = asyncio.new_event_loop()
                            try:
                                return _loop.run_until_complete(
                                    self._assistant.handle_message(user_message, user_id)
                                )
                            finally:
                                _loop.close()
    
                        result = await loop.run_in_executor(None, _sync_handle)
                        logger.info(f"[DEBUG] handle_message 完成, 耗时 {_time.time()-_t0:.1f}s")
    
                        reply_text = result.get("text", "")
                        files = result.get("files", [])
                        logger.info(f"[DEBUG] 回复文本长度: {len(reply_text)}, 文件数: {len(files)}")
    
                        file_server = self._config.get("dingtalk_bot", {}).get(
                            "file_server_url", "http://127.0.0.1:5679")
    
                        # 如果有文件，附加下载链接
                        file_links = ""
                        if files:
                            file_links += "\n\n---\n📥 **文件下载:**\n"
                            for f in files:
                                url = f.get("url", "")
                                if url and not url.startswith("http"):
                                    url = f"{file_server}{url}"
                                file_links += f"- [{f['name']}]({url})\n"
    
                        # 钉钉 Markdown 消息字符限制约 5000 字
                        DINGTALK_MAX_LEN = 4500
    
                        if len(reply_text) + len(file_links) <= DINGTALK_MAX_LEN:
                            # 短回复: 直接发送
                            logger.info(f"[DEBUG] 发送 Markdown 回复 ({len(reply_text)} chars)...")
                            resp = self.reply_markdown("AI 助理回复", reply_text + file_links,
                                                       incoming_message)
                            logger.info(f"[DEBUG] Markdown 回复结果: {resp}")
                        else:
                            # 长回复: 生成摘要 + 保存完整内容为 Word 附件
                            logger.info(f"[DEBUG] 回复过长({len(reply_text)} chars), 生成摘要+Word附件")
    
                            # 保存完整内容为 Word 文档
                            word_url = self._save_as_word(reply_text, user_message)
                            if word_url:
                                download_link = f"{file_server}{word_url}"
                                file_links += f"\n📄 **完整结果:** [{os.path.basename(word_url)}]({download_link})\n"
    
                            # 生成简要摘要（取前几条 + 统计信息）
                            summary = self._make_summary(reply_text, DINGTALK_MAX_LEN - len(file_links) - 200)
                            final_text = summary + file_links
    
                            logger.info(f"[DEBUG] 发送摘要 ({len(final_text)} chars) + Word附件")
                            resp = self.reply_markdown("AI 助理回复", final_text, incoming_message)
                            logger.info(f"[DEBUG] 摘要回复结果: {resp}")

                    except Exception as e:
                        logger.error(f"处理消息失败: {e}", exc_info=True)
                        try:
                            self.reply_text(f"❌ 处理消息时发生错误: {str(e)[:200]}", incoming_message)
                        except Exception as e2:
                            logger.error(f"发送错误回复也失败: {e2}", exc_info=True)

                # 启动后台任务
                asyncio.create_task(_background_process())
                
                # 立即返回 OK，不让钉钉等待
                return AckMessage.STATUS_OK, "OK"


            def _save_as_word(self, text: str, query: str) -> str:
                """将完整回复内容保存为 Word 文档，返回下载路径"""
                try:
                    from docx import Document
                    from docx.shared import Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    from datetime import datetime as dt
                except ImportError:
                    logger.warning("python-docx 未安装，无法生成 Word 附件")
                    # 回退到 txt 文件
                    return self._save_as_txt(text, query)

                try:
                    ts = dt.now().strftime("%Y%m%d_%H%M%S")
                    safe_query = "".join(c for c in query[:20] if c.isalnum() or c in " _-\u4e00-\u9fff")
                    filename = f"{safe_query}_{ts}.docx"
                    filepath = os.path.join(EXPORTS_DIR, filename)
                    os.makedirs(EXPORTS_DIR, exist_ok=True)

                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = '宋体'
                    style.font.size = Pt(10.5)

                    # 标题
                    title = doc.add_heading(f'AI 研究助理 - 文献调研结果', level=0)
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

                    info = doc.add_paragraph()
                    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = info.add_run(f'查询: {query}  |  生成时间: {dt.now().strftime("%Y-%m-%d %H:%M")}')
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(128, 128, 128)

                    doc.add_paragraph()

                    # 解析 Markdown 写入 Word
                    for line in text.split('\n'):
                        line_stripped = line.strip()
                        if not line_stripped:
                            doc.add_paragraph()
                            continue
                        if line_stripped.startswith('## '):
                            doc.add_heading(line_stripped[3:], level=2)
                        elif line_stripped.startswith('### '):
                            doc.add_heading(line_stripped[4:], level=3)
                        elif line_stripped.startswith('# '):
                            doc.add_heading(line_stripped[2:], level=1)
                        elif line_stripped.startswith('**') and line_stripped.endswith('**'):
                            p = doc.add_paragraph()
                            run = p.add_run(line_stripped.strip('*'))
                            run.bold = True
                        elif line_stripped.startswith('> '):
                            p = doc.add_paragraph(line_stripped[2:])
                            p.paragraph_format.left_indent = Pt(36)
                            for run in p.runs:
                                run.font.size = Pt(9)
                                run.font.color.rgb = RGBColor(100, 100, 100)
                        elif line_stripped.startswith('- ') or line_stripped.startswith('* '):
                            doc.add_paragraph(line_stripped[2:], style='List Bullet')
                        else:
                            # 处理加粗文本
                            if '**' in line_stripped:
                                p = doc.add_paragraph()
                                parts = line_stripped.split('**')
                                for i, part in enumerate(parts):
                                    if part:
                                        run = p.add_run(part)
                                        if i % 2 == 1:
                                            run.bold = True
                            else:
                                doc.add_paragraph(line_stripped)

                    doc.save(filepath)
                    logger.info(f"完整结果已保存为 Word: {filepath}")
                    return f"/download/{filename}"
                except Exception as e:
                    logger.error(f"保存 Word 失败: {e}", exc_info=True)
                    return self._save_as_txt(text, query)

            def _save_as_txt(self, text: str, query: str) -> str:
                """回退: 保存为纯文本文件"""
                try:
                    from datetime import datetime as dt
                    ts = dt.now().strftime("%Y%m%d_%H%M%S")
                    safe_query = "".join(c for c in query[:20] if c.isalnum() or c in " _-\u4e00-\u9fff")
                    filename = f"{safe_query}_{ts}.txt"
                    filepath = os.path.join(EXPORTS_DIR, filename)
                    os.makedirs(EXPORTS_DIR, exist_ok=True)

                    # 清除 Markdown 标记用于纯文本
                    clean_text = text.replace('**', '').replace('> ', '  ')
                    with open(filepath, 'w', encoding='utf-8') as f:
                        f.write(f"查询: {query}\n")
                        f.write(f"生成时间: {dt.now().strftime('%Y-%m-%d %H:%M')}\n")
                        f.write("=" * 60 + "\n\n")
                        f.write(clean_text)

                    logger.info(f"完整结果已保存为文本: {filepath}")
                    return f"/download/{filename}"
                except Exception as e:
                    logger.error(f"保存文本文件也失败: {e}")
                    return ""

            @staticmethod
            def _make_summary(text: str, max_len: int) -> str:
                """从完整回复中提取摘要（标题 + 前几条结果 + 统计）"""
                lines = text.split('\n')
                summary_lines = []
                paper_count = 0
                total_papers = 0
                in_header = True
                max_show = 5  # 摘要中最多显示几条文献

                for line in lines:
                    stripped = line.strip()

                    # 保留头部信息（标题、搜索说明等）
                    if in_header:
                        summary_lines.append(line)
                        if stripped == '':
                            # 检查是否已经过了头部区域
                            if any(s.startswith('## ') or s.startswith('> 找到') for s in summary_lines):
                                in_header = False
                        continue

                    # 计算论文总数
                    if stripped.startswith('**') and stripped[2:3].isdigit():
                        total_papers += 1

                    # 摘要中只保留前 N 条文献
                    if total_papers <= max_show:
                        summary_lines.append(line)
                    elif total_papers == max_show + 1 and paper_count == 0:
                        paper_count = 1  # 标记已超出

                result = '\n'.join(summary_lines)

                # 如果超出限制，按行截断
                if len(result) > max_len:
                    result = result[:max_len]
                    last_newline = result.rfind('\n')
                    if last_newline > max_len // 2:
                        result = result[:last_newline]

                # 添加省略提示
                if total_papers > max_show:
                    result += f"\n\n---\n> ⚡ 以上仅显示前 {max_show} 条，共 {total_papers} 篇文献。完整结果请查看附件。\n"

                return result

        return _Handler(assistant, config)


# ============================================================
# Webhook-only 模式（仅推送，不接收消息）
# 适用于只有自定义 Webhook 机器人的情况
# ============================================================

class DingTalkBotWebhookOnly:
    """
    仅 Webhook 推送模式

    当用户还没有配置企业内部应用时，可以先用这个模式。
    提供 HTTP 接口，用户可以通过浏览器/curl 发送查询。
    """

    def __init__(self, config: dict):
        self.config = config
        self.assistant = ResearchAssistant(config)
        self.app = Flask(__name__)
        self._setup_routes()

    def _setup_routes(self):
        @self.app.route("/")
        def index():
            return """
            <html><head><title>AI 研究助理</title></head>
            <body style="font-family: sans-serif; max-width: 800px; margin: 50px auto; padding: 20px;">
            <h1>🤖 AI 研究助理</h1>
            <p>使用 API 接口发送查询:</p>
            <pre>POST /api/ask
Content-Type: application/json
{"message": "搜索机器人抓取相关的论文"}</pre>
            <p>或直接在下方输入:</p>
            <form id="askForm">
                <textarea id="msg" rows="3" style="width:100%;font-size:16px;" placeholder="输入你的问题..."></textarea>
                <button type="submit" style="margin-top:10px;padding:10px 30px;font-size:16px;">提交</button>
            </form>
            <div id="result" style="margin-top:20px;white-space:pre-wrap;"></div>
            <script>
            document.getElementById('askForm').onsubmit = async (e) => {
                e.preventDefault();
                const msg = document.getElementById('msg').value;
                document.getElementById('result').textContent = '处理中...';
                const resp = await fetch('/api/ask', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({message: msg})
                });
                const data = await resp.json();
                document.getElementById('result').textContent = data.text || JSON.stringify(data);
            };
            </script>
            </body></html>
            """

        @self.app.route("/api/ask", methods=["POST"])
        def api_ask():
            from flask import request
            data = json.loads(request.data) if request.data else {}
            message = data.get("message", "")
            if not message:
                return jsonify({"error": "请提供 message 参数"})

            import asyncio
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            try:
                result = loop.run_until_complete(
                    self.assistant.handle_message(message)
                )
                return jsonify(result)
            except Exception as e:
                return jsonify({"error": str(e)})
            finally:
                loop.close()

    def start(self, host: str = "0.0.0.0", port: int = 5680):
        logger.info(f"🤖 AI 研究助理（Web 模式）启动: http://{host}:{port}")
        self.app.run(host=host, port=port, debug=False)


# ============================================================
# 入口
# ============================================================

def main():
    parser = argparse.ArgumentParser(description="钉钉 AI 助理机器人")
    parser.add_argument("--port", type=int, default=5679,
                        help="文件下载服务端口 (默认: 5679)")
    parser.add_argument("--webhook-only", action="store_true",
                        help="仅使用 Webhook 推送模式（提供 Web 界面，不需要企业应用）")
    parser.add_argument("--web-port", type=int, default=5680,
                        help="Web 模式端口 (默认: 5680)")
    parser.add_argument("--config", type=str, default=None,
                        help="配置文件路径")
    args = parser.parse_args()

    # 加载配置
    try:
        config = load_config(args.config)
    except Exception as e:
        print(f"❌ 配置文件加载失败: {e}")
        sys.exit(1)

    os.makedirs(EXPORTS_DIR, exist_ok=True)

    # 设置 stdout/stderr 编码（Windows GBK 兼容）
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
        sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

    mode_str = "Webhook (Web界面)" if args.webhook_only else "Stream (钉钉长连接)"
    print(f"""
========================================
  DingTalk AI Research Assistant
========================================
  File Server : http://0.0.0.0:{args.port}
  Mode        : {mode_str}
========================================
    """)

    # 启动文件下载服务（后台线程）
    file_thread = threading.Thread(
        target=start_file_server,
        args=(args.port,),
        daemon=True
    )
    file_thread.start()
    logger.info(f"📂 文件下载服务已启动: http://0.0.0.0:{args.port}")

    if args.webhook_only:
        # Web 模式
        bot = DingTalkBotWebhookOnly(config)
        bot.start(port=args.web_port)
    else:
        # Stream 模式
        try:
            bot = DingTalkBotStream(config)
            bot.start()
        except ValueError as e:
            logger.error(str(e))
            logger.info("\n💡 提示: 如果还没有配置企业内部应用，可以先用 --webhook-only 模式:")
            logger.info("   python dingtalk_bot_server.py --webhook-only")
            sys.exit(1)


if __name__ == "__main__":
    main()
